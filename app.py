import math
from flask import Flask, request, jsonify
import re
import string
import numpy as np
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
import nltk
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk.tokenize import sent_tokenize
import PyPDF2
import docx
import flask_cors
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

nltk.download("punkt")
nltk.download("stopwords")
nltk.download('wordnet')
nltk.download('punkt_tab')
from scraper import scrape_website

app = Flask(__name__)

# Custom stopwords
english_stopset = list(set(stopwords.words('english')))

lemmatizer = WordNetLemmatizer()

# Helper function to extract text from a PDF file
def extract_text_from_pdf(pdf_file):
    try:
        reader = PyPDF2.PdfReader(pdf_file)
        
        logger.info(f"Number of pages: {len(reader.pages)}")
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ''  # Handle None if no text is extracted
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return ""

# Helper function to extract text from a Word file
def extract_text_from_word(word_file):
    try:
        doc = docx.Document(word_file)
        logger.info(f"Number of paragraphs: {len(doc.paragraphs)}")
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        logger.error(f"Error extracting text from Word file: {e}")
        return ""

# Clean and lemmatize text
def clean_and_lemmatize(text):
    cleaned_sentences = []
    sentences = sent_tokenize(text)
    for sent in sentences:
        sent = re.sub(r'[^\x00-\x7F]+', ' ', sent)  # Replace non-ASCII characters
        sent = re.sub(r'@\w+', '', sent)  # Remove mentions
        sent = sent.lower()  # Lowercase
        sent = re.sub(r'[%s]' % re.escape(string.punctuation), ' ', sent)  # Remove punctuation
        sent = re.sub(r'[0-9]', '', sent)  # Remove numbers
        sent = re.sub(r'\s{2,}', ' ', sent)  # Remove extra spaces
        lemmatized = ' '.join([lemmatizer.lemmatize(word) for word in nltk.word_tokenize(sent)])
        if len(lemmatized.split()) > 3:  # Only keep sentences with more than 3 words
            cleaned_sentences.append(lemmatized)
    return cleaned_sentences

# Vectorize the sentences
def vectorize_sentences(cleaned_sentences):
    vectorizer = TfidfVectorizer(
            analyzer='word',
            ngram_range=(1, 2),
            stop_words=english_stopset,
            lowercase=True
    )
    X = vectorizer.fit_transform(cleaned_sentences)
    return vectorizer, pd.DataFrame(X.T.toarray())

# Function to get similar sentences
def get_similar_sentences(query, df, cleaned_sentences, vectorizer):
    q_vec = vectorizer.transform([query]).toarray().reshape(df.shape[0],)
    sim = {}
    
    for i in range(len(cleaned_sentences)):
        norm_df = np.linalg.norm(df.loc[:, i])
        norm_q_vec = np.linalg.norm(q_vec)
        if norm_df == 0 or norm_q_vec == 0:
            sim[i] = 0
        else:
            sim[i] = np.dot(df.loc[:, i].values, q_vec) / (norm_df * norm_q_vec)
    
    # Sort by similarity and return top 5 sentences
    sim_sorted = sorted(sim.items(), key=lambda x: x[1], reverse=True)[:5]
    results = [(cleaned_sentences[i], sim[i]) for i, _ in sim_sorted if sim[i] > 0]
    
    return results

# Flask route to handle the file upload and search
@app.route('/search', methods=['POST'])
@flask_cors.cross_origin()
def search():
    data = request.get_json()
    logger.info(f"Received data: {data}")
    query = data.get('q')
    siteUrl = data.get('url')
   
    if not query:
        return jsonify({'error': 'Query is required'}), 400

    scraped_contents = scrape_website(siteUrl)
    logger.info(f"Scraped {len(scraped_contents)} pages")

    all_results = []
    for url, text in scraped_contents.items():
        logger.info(f"Processing URL: {url}")
        
        # Clean and lemmatize the extracted text
        cleaned_sentences = clean_and_lemmatize(text)
        logger.info(f"Number of cleaned sentences: {len(cleaned_sentences)}")
        
        if not cleaned_sentences:
            logger.warning(f"No valid sentences found for URL: {url}")
            continue
        
        # Vectorize the sentences
        try:
            vectorizer, df = vectorize_sentences(cleaned_sentences)
        except Exception as e:
            logger.error(f"Error vectorizing sentences for URL {url}: {e}")
            continue

        # Clean and lemmatize the query
        query_cleaned = ' '.join([lemmatizer.lemmatize(word) for word in nltk.word_tokenize(query.lower())])
        
        # Get similar sentences
        try:
            searchResult = get_similar_sentences(query_cleaned, df, cleaned_sentences, vectorizer)
            logger.info(f"Search results for {url}: {len(searchResult)}")
        except Exception as e:
            logger.error(f"Error getting similar sentences for URL {url}: {e}")
            continue

        # Filter out NaN scores and create result objects
        url_results = [
            {"url": url, "sentence": sentence, "score": score}
            for sentence, score in searchResult
            if not math.isnan(score)
        ]
        logger.info(f"Final results for {url}: {len(url_results)}")
        all_results.extend(url_results)

    # Sort all results by score
    all_results.sort(key=lambda x: x['score'], reverse=True)
    
    # Take top 10 results
    top_results = all_results[:10]
    
    logger.info(f"Returning {len(top_results)} results")
    return jsonify(top_results)

# Run the Flask app
if __name__ == "__main__":
    app.run(debug=True)